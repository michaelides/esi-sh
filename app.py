
import os
import json # Used for HF persistence
import re
import uuid
import shutil # For file operations (copying uploaded files)
from typing import List, Dict, Any, Tuple, Optional
from io import BytesIO # For creating in-memory byte streams (e.g., for DOCX downloads)

import pandas as pd
from pypdf import PdfReader
from docx import Document as DocxDocument # python-docx library
try:
    import pyreadstat # For SPSS .sav files
except ImportError:
    pyreadstat = None
    print("Warning: pyreadstat not installed. SPSS (.sav) file support will be disabled.")

from shiny import App, ui, render, reactive, Session
from dotenv import load_dotenv
from huggingface_hub import HfFileSystem, hf_hub_download # For HF persistence


from llama_index.core.llms import ChatMessage, MessageRole
from llama_index.core import Settings
from llama_index.core.tools import FunctionTool

# Local module imports
from agent import (
    create_orchestrator_agent,
    initialize_settings as initialize_agent_settings,
    DEFAULT_PROMPTS,
    generate_llm_greeting as agent_generate_llm_greeting,
)
from tools import UI_ACCESSIBLE_WORKSPACE # Path to where uploaded files are temporarily stored
from config import HF_USER_MEMORIES_DATASET_ID # HF Dataset ID for storing user chat data
from shiny_ui import get_app_ui, chat_options_modal_ui, delete_chat_confirmation_modal_ui

# --- Global Setup & Constants ---
PROJECT_ROOT = os.path.abspath(os.path.dirname(__file__))
os.makedirs(UI_ACCESSIBLE_WORKSPACE, exist_ok=True) # Ensure workspace directory exists

MAX_CHAT_HISTORY_MESSAGES = 15 # Max messages to keep in agent's short-term history
load_dotenv() # Load environment variables from .env file
fs = HfFileSystem() # Global Hugging Face FileSystem client

# Global flags and variables for LLM and Agent initialization status
llm_settings_initialized: bool = False
llm_settings_error: str | None = None
orchestrator_agent_rv = reactive.Value(None) # Holds the agent instance, reactive for UI updates
agent_init_error_rv = reactive.Value(None)   # Holds agent initialization error, reactive

# Dictionary to allow global tool functions to access server-specific reactive values
# This is a workaround for FunctionTool not being able to directly access Shiny's reactive scope easily.
app_reactive_state_access = {
    "get_uploaded_docs": None,
    "get_uploaded_dfs": None,
}

# --- Core Application Logic Functions ---

def setup_global_llm_settings() -> Tuple[bool, str | None]:
    """Initializes global LlamaIndex settings for LLM and embedding models."""

    try:
        initialize_agent_settings() # From agent.py
        print("LLM settings initialized successfully.")
        return True, None
    except Exception as e:

        print(f"Fatal Error: Could not initialize LLM settings. {e}")
        return False, f"Fatal Error: Could not initialize LLM settings. {e}"

# Initialize LLM settings once at application startup
llm_settings_initialized, llm_settings_error = setup_global_llm_settings()

def setup_agent_shiny(
    rv_docs_ref: reactive.Value,
    rv_dfs_ref: reactive.Value,
    max_search_results: int = 10
) -> Tuple[Any | None, str | None]:
    """
    Initializes or re-initializes the LlamaIndex agent.
    Tool functions (actual_read_doc_fn, actual_analyze_df_fn) are defined *inside* this
    function to create closures over the reactive value references (rv_docs_ref, rv_dfs_ref),
    allowing them to access the current state of uploaded files when the agent executes them.
    """
    if not llm_settings_initialized:
        return None, "LLM settings must be initialized before agent setup."

    def actual_read_doc_fn(filename: str) -> str:
        """Tool function: Reads content of an uploaded document."""
        docs = rv_docs_ref.get() # Access reactive value for documents
        return docs.get(filename, f"Error: Document '{filename}' not found. Available: {list(docs.keys())}")

    def actual_analyze_df_fn(filename: str, head_rows: int = 5) -> str:
        """Tool function: Analyzes an uploaded DataFrame."""
        dfs = rv_dfs_ref.get() # Access reactive value for dataframes
        if filename not in dfs:
            return f"Error: DataFrame '{filename}' not found. Available: {list(dfs.keys())}"
        df = dfs[filename]
        if not isinstance(df, pd.DataFrame):
            return f"Error: '{filename}' is not a pandas DataFrame object."

        # Construct summary string
        info = f"DataFrame: {filename}\nShape: {df.shape}\nColumns: {list(df.columns)}\nData Types:\n{df.dtypes.to_string()}\n"
        if head_rows > 0:
            info += f"First {min(head_rows, len(df))} rows:\n{df.head(min(head_rows, len(df))).to_string()}\n"
        info += f"Summary Statistics:\n{df.describe(include='all').to_string()}" # include='all' for mixed types
        return info

    try:
        dynamic_tools = [
            FunctionTool.from_defaults(fn=actual_read_doc_fn, name="read_uploaded_document", description="Reads text content from a previously uploaded document (PDF, DOCX, TXT, MD)."),
            FunctionTool.from_defaults(fn=actual_analyze_df_fn, name="analyze_uploaded_dataframe", description="Provides a summary analysis of a previously uploaded tabular dataset (CSV, XLSX, SAV).")
        ]
        agent_instance = create_orchestrator_agent(dynamic_tools=dynamic_tools, max_search_results=max_search_results)
        print(f"Agent (re)initialized with max_search_results: {max_search_results}.")
        return agent_instance, None
    except Exception as e:
        print(f"Error during agent (re)initialization: {e}")
        return None, f"Failed to (re)initialize agent: {e}"

def generate_suggested_prompts_shiny(chat_history: List[Dict[str, Any]]) -> List[str]:
    """Generates suggested prompts. Currently returns defaults."""
    # TODO: Consider enabling agent_generate_suggested_prompts(chat_history) from agent.py if performance is acceptable.
    # print(f"generate_suggested_prompts_shiny called. History length: {len(chat_history)}. Returning defaults.")
    return list(DEFAULT_PROMPTS)

def format_chat_history_shiny(messages: List[Dict[str, Any]]) -> List[ChatMessage]:
    """Converts app's message format to LlamaIndex's ChatMessage format."""
    return [ChatMessage(role=MessageRole.USER if m["role"] == "user" else MessageRole.ASSISTANT, content=m["content"])
            for m in messages[-MAX_CHAT_HISTORY_MESSAGES:]]

def get_agent_response_shiny(query: str, chat_history: List[ChatMessage], temp: float, verbosity: int, search_results: int) -> str:
    """Gets a response from the LlamaIndex agent."""
    agent_instance = orchestrator_agent_rv.get()
    if not agent_instance:
        return "Error: AI agent is not available or not yet initialized."
    try:
        if Settings.llm and hasattr(Settings.llm, 'temperature'):
            Settings.llm.temperature = temp # Update LLM temperature

        # Verbosity is prepended to the query for the agent to handle.
        # search_results are handled by re-initializing the agent if the count changes.
        modified_query = f"Verbosity Level: {verbosity}. {query}"
        response = agent_instance.chat(modified_query, chat_history=chat_history)
        return response.response if hasattr(response, 'response') else str(response)
    except Exception as e:
        print(f"Error during agent chat: {e}")
        return f"Agent error: {e}"

def process_uploaded_file_shiny(file_info: dict, rv_docs: reactive.Value, rv_dfs: reactive.Value) -> Tuple[Optional[str], Optional[str]]:
    """Processes an uploaded file, extracts content/data, and updates reactive values."""
    file_name, temp_file_path = file_info['name'], file_info['tempfile']
    ext = os.path.splitext(file_name)[1].lower()
    ws_path = os.path.join(UI_ACCESSIBLE_WORKSPACE, file_name) # Path in shared workspace

    try:
        shutil.copy(temp_file_path, ws_path) # Copy to workspace
        content, file_type_processed = None, None

        if ext == ".pdf":
            content = "".join(p.extract_text() for p in PdfReader(ws_path).pages if p.extract_text())
            file_type_processed = "document"
        elif ext == ".docx":
            content = "\n".join([p.text for p in DocxDocument(ws_path).paragraphs])
            file_type_processed = "document"
        elif ext in [".txt", ".md"]:
            with open(ws_path, "r", encoding="utf-8") as f: content = f.read()
            file_type_processed = "document"
        elif ext == ".csv":
            content, file_type_processed = pd.read_csv(ws_path), "dataframe"
        elif ext == ".xlsx":
            content, file_type_processed = pd.read_excel(ws_path), "dataframe"
        elif ext == ".sav" and pyreadstat:
            content, file_type_processed = pyreadstat.read_sav(ws_path)[0], "dataframe" # [0] is the DataFrame

        if file_type_processed == "document":
            docs_copy = rv_docs.get().copy(); docs_copy[file_name] = content; rv_docs.set(docs_copy)
        elif file_type_processed == "dataframe":
            dfs_copy = rv_dfs.get().copy(); dfs_copy[file_name] = content; rv_dfs.set(dfs_copy)
        else:
            print(f"Unsupported file type: {ext} for file {file_name}")
            return None, file_name # Type not processed, but file is in workspace

        print(f"File '{file_name}' processed as {file_type_processed}.")
        return file_type_processed, file_name
    except Exception as e:
        print(f"Error processing uploaded file '{file_name}': {e}")
        return None, file_name # Error during processing

def get_discussion_markdown_shiny(chat_id: str, all_msgs: dict, meta: dict) -> str:
    """Generates markdown content for a given chat session."""
    messages = all_msgs.get(chat_id, [])
    chat_name = meta.get(chat_id, "Unknown Chat")
    markdown_content = [f"# Chat Discussion: {chat_name}\n\n"]
    for msg in messages:
        role = msg.get("role", "unknown").capitalize()
        content = msg.get("content", "")
        markdown_content.append(f"**{role}:**\n{content}\n\n---\n")
    return "".join(markdown_content)

def get_discussion_docx_shiny(chat_id: str, all_msgs: dict, meta: dict) -> bytes:
    """Generates DOCX file content for a given chat session."""
    messages = all_msgs.get(chat_id, [])
    chat_name = meta.get(chat_id, "Unknown Chat")
    doc = DocxDocument()
    doc.add_heading(f"Chat Discussion: {chat_name}", level=1)
    for msg in messages:
        role = msg.get("role", "unknown").capitalize()
        content = msg.get("content", "")
        doc.add_heading(f"{role}:", level=3)
        doc.add_paragraph(content)
        doc.add_paragraph("---")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

# --- Hugging Face Persistence Functions ---
def load_user_data_from_hf_shiny(user_id: str, rv_meta: reactive.Value, rv_all_msgs: reactive.Value):
    """Loads user's chat metadata and all messages from Hugging Face Hub."""
    hf_token = os.getenv("HF_TOKEN")
    if not hf_token:
        print("HF_TOKEN not set. Cannot load user data from Hugging Face.")
        rv_meta.set({}); rv_all_msgs.set({}) # Ensure clean state if load fails
        return

    meta_fn = f"user_memories/{user_id}_metadata.json"
    msgs_fn = f"user_memories/{user_id}_messages.json"
    loaded_meta, loaded_msgs = {}, {}

    try:
        meta_path = hf_hub_download(HF_USER_MEMORIES_DATASET_ID, meta_fn, repo_type="dataset", token=hf_token)
        with open(meta_path, "r") as f: loaded_meta = json.load(f)
        print(f"Metadata loaded for user {user_id}.")
    except Exception as e:
        print(f"Metadata file not found or error for user {user_id} ({meta_fn}): {e}. Initializing empty metadata.")


# --- Core Logic Functions (Adapted for Shiny) ---
def setup_global_llm_settings() -> Tuple[bool, str | None]:
    print("Attempting to initialize LLM settings...")
    try:

        msgs_path = hf_hub_download(HF_USER_MEMORIES_DATASET_ID, msgs_fn, repo_type="dataset", token=hf_token)
        with open(msgs_path, "r") as f: loaded_msgs = json.load(f)
        print(f"All messages loaded for user {user_id}.")
    except Exception as e:
        print(f"Messages file not found or error for user {user_id} ({msgs_fn}): {e}. Initializing empty messages.")
        
    rv_meta.set(loaded_meta)
    rv_all_msgs.set(loaded_msgs)

def save_chat_metadata_shiny(user_id: str, chat_meta_data: Dict):
    """Saves user's chat metadata to Hugging Face Hub."""
    if not user_id: print("Error: No user_id provided for saving chat metadata."); return
    hf_token = os.getenv("HF_TOKEN")
    if not hf_token: print("HF_TOKEN not set. Cannot save metadata to Hugging Face."); return

    fn_local = f"{user_id}_metadata.json" # Temporary local file
    path_in_repo = f"user_memories/{fn_local}" # Path in the HF dataset repository

    try:
        with open(fn_local, "w") as f: json.dump(chat_meta_data, f, indent=2)
        fs.upload_file(fn_local, path_in_repo, HF_USER_MEMORIES_DATASET_ID, repo_type="dataset", token=hf_token, commit_message=f"Update chat metadata for user {user_id}")
        print(f"Chat metadata saved to HF for user {user_id}.")
        os.remove(fn_local) # Clean up local temp file
    except Exception as e:
        print(f"Error saving chat metadata to Hugging Face for user {user_id}: {e}")

def save_chat_history_shiny(user_id: str,
                            updated_chat_id: Optional[str] = None,
                            updated_messages_list: Optional[List[Dict]] = None,
                            messages_dict_to_save_directly: Optional[Dict] = None):
    """Saves all of a user's chat messages to a single JSON file on Hugging Face Hub."""
    if not user_id: print("Error: No user_id provided for saving chat history."); return
    hf_token = os.getenv("HF_TOKEN")
    if not hf_token: print("HF_TOKEN not set. Cannot save chat history to Hugging Face."); return

    messages_filename_local = f"{user_id}_messages.json"
    repo_path_in_dataset = f"user_memories/{messages_filename_local}"
    all_user_messages_data = {}

    if messages_dict_to_save_directly is not None:
        # Used when the entire message blob is being overwritten (e.g., after deleting a chat)
        all_user_messages_data = messages_dict_to_save_directly
        # print(f"Saving all messages directly for user {user_id}.") # Can be verbose
    else:
        # Standard update: load existing, modify one chat, save all
        try:
            msgs_file_path = hf_hub_download(
                repo_id=HF_USER_MEMORIES_DATASET_ID,
                filename=repo_path_in_dataset,
                repo_type="dataset",
                token=hf_token
            )
            with open(msgs_file_path, "r") as f: all_user_messages_data = json.load(f)
            # print(f"Existing messages data loaded for user {user_id}.") # Can be verbose
        except Exception as e:
            print(f"No existing messages file or error for user {user_id} (Path: {repo_path_in_dataset}): {e}. Starting fresh.")
            all_user_messages_data = {} # Initialize if not found or error

        if updated_chat_id and updated_messages_list is not None:
            all_user_messages_data[updated_chat_id] = updated_messages_list
            # print(f"Updated messages for chat {updated_chat_id} for user {user_id}.") # Can be verbose
        elif updated_chat_id and updated_messages_list is None:
             print(f"Warning: updated_chat_id provided for user {user_id} but no messages list for update.")

    try:
        with open(messages_filename_local, "w") as f: json.dump(all_user_messages_data, f, indent=2)
        fs.upload_file(
            path_or_fileobj=messages_filename_local,
            path_in_repo=repo_path_in_dataset,
            repo_id=HF_USER_MEMORIES_DATASET_ID,
            repo_type="dataset",
            token=hf_token,
            commit_message=f"Update chat messages for user {user_id}"
        )
        print(f"Chat messages saved to HF for user {user_id}.")
        os.remove(messages_filename_local)
    except Exception as e:
        print(f"Error saving chat messages to Hugging Face for user {user_id}: {e}")

# --- App Entry Point & Server Function ---
if not os.getenv("GOOGLE_API_KEY"): print("Warning: GOOGLE_API_KEY environment variable not set.")
app_ui = get_app_ui()

def server(input: reactive.Input, output: reactive.Output, session: Session):
    print(f"--- Server Function Initializing for session {session.id} ---")
    
    # --- Reactive Values Initialization ---
    rv_user_id = reactive.Value(str(uuid.uuid4()))
    rv_long_term_memory_enabled = reactive.Value(True)
    rv_chat_metadata = reactive.Value({})
    rv_all_chat_messages = reactive.Value({})
    rv_current_chat_id = reactive.Value(None)
    rv_messages = reactive.Value([])
    rv_chat_modified = reactive.Value(False)
    rv_suggested_prompts = reactive.Value(list(DEFAULT_PROMPTS))
    rv_llm_temperature = reactive.Value(0.7)
    rv_llm_verbosity = reactive.Value(3)
    rv_search_results_count = reactive.Value(10)
    rv_uploaded_documents = reactive.Value({})
    rv_uploaded_dataframes = reactive.Value({})
    
    # For FunctionTools to access reactive state
    app_reactive_state_access["get_uploaded_docs"] = rv_uploaded_documents.get
    app_reactive_state_access["get_uploaded_dfs"] = rv_uploaded_dataframes.get
    
    # For dynamic button click detection
    last_action_button_values = reactive.Value({})
    # For clipboard functionality
    rv_clipboard_text = reactive.Value(None)
    # For chat options modal
    rv_options_modal_chat_id = reactive.Value(None)
    rv_show_delete_confirmation_modal = reactive.Value(False)

    # Flags for sequencing cookie processing and initial setup
    _cookie_ltm_processed = reactive.Value(False)
    _cookie_user_id_processed = reactive.Value(False)
    _initial_chat_selected_after_load = reactive.Value(False) # Ensures initial chat selection runs once

    # --- Cookie Handling and Initial Setup Effects ---
    # This sequence of effects handles loading LTM preference and user ID from cookies,
    # then initializes the agent and the first chat session.

    @reactive.Effect
    @reactive.event(input.cookie_ltm_preference_loaded)
    def handle_ltm_cookie_load():
        """Handles the LTM preference cookie value received from JavaScript."""
        cookie_val = input.cookie_ltm_preference_loaded()
        # print(f"LTM preference cookie loaded from JS: '{cookie_val}'") # Debug
        if cookie_val and cookie_val != "None":
            current_ltm_state = cookie_val == "true"
            rv_long_term_memory_enabled.set(current_ltm_state)
            ui.update_switch("long_term_memory_enabled", value=current_ltm_state) # Sync UI
        else: # No cookie, default to True and tell JS to set it
            session.send_custom_message("setCookie", {"name": "ltm_preference", "value": "true", "days": 365})
            rv_long_term_memory_enabled.set(True)
            ui.update_switch("long_term_memory_enabled", value=True)
        _cookie_ltm_processed.set(True)

    @reactive.Effect
    @reactive.event(input.cookie_user_id_loaded)
    def handle_user_id_cookie_load():
        """Handles the User ID cookie value received from JavaScript. Depends on LTM pref being known."""
        if not _cookie_ltm_processed.get(): return # Wait for LTM pref to be processed

        cookie_val = input.cookie_user_id_loaded()
        # print(f"User ID cookie loaded from JS: '{cookie_val}'") # Debug
        current_ltm_enabled = rv_long_term_memory_enabled.get()

        if cookie_val and cookie_val != "None": # Cookie exists
            if current_ltm_enabled:
                rv_user_id.set(cookie_val) # Use existing cookie ID if LTM is on
            else: # LTM off, but cookie exists: tell JS to delete it. rv_user_id remains the initial temp one.
                session.send_custom_message("deleteCookie", {"name": "user_id"})
        # else: No cookie. If LTM is on, ensure_persistent_user_id_and_load_data will create one.
        # If LTM is off, rv_user_id is already a temp one.

        _cookie_user_id_processed.set(True)
        # This is a critical point: now that LTM pref and User ID cookie (or lack thereof) are known,
        # finalize user ID and load their data if applicable.
        ensure_persistent_user_id_and_load_data()


    @reactive.Effect
    @reactive.event(input.long_term_memory_enabled, ignore_init=True)
    def handle_ltm_toggle_interaction():
        """Handles manual toggle of the LTM switch by the user *after* initial cookie load."""
        if not (_cookie_ltm_processed.get() and _cookie_user_id_processed.get()): return # Only after initial setup

        new_ltm_state = input.long_term_memory_enabled()
        rv_long_term_memory_enabled.set(new_ltm_state) # Update reactive value
        session.send_custom_message("setCookie", {"name": "ltm_preference", "value": str(new_ltm_state).lower(), "days": 365})
        print(f"LTM UI toggle changed to: {new_ltm_state}, cookie updated.")

        if not new_ltm_state: # LTM was turned OFF
            session.send_custom_message("deleteCookie", {"name": "user_id"}) # Delete persistent ID cookie
            rv_user_id.set(str(uuid.uuid4())) # Reset to a new temporary ID
            rv_chat_metadata.set({}); rv_all_chat_messages.set({}) # Clear local LTM data
            _initial_chat_selected_after_load.set(False) # Allow re-selection for new temp session
            create_new_chat_session_shiny(force_temp_name=True)
        else: # LTM was turned ON
            # This will ensure a persistent ID is set and attempt to load data
            ensure_persistent_user_id_and_load_data()
            # ensure_persistent_user_id_and_load_data calls load_user_data, which will trigger handle_initial_chat_selection
            # If no data loaded, handle_initial_chat_selection will call create_new_chat_session_shiny.

    def ensure_persistent_user_id_and_load_data():
        """Ensures user_id is persistent if LTM is ON, sets cookie, and loads data."""
        # This function is called after LTM preference is known (either from cookie or toggle).
        if rv_long_term_memory_enabled.get():
            current_user_id = rv_user_id.get() # Could be temp or from cookie
            cookie_user_id = input.cookie_user_id_loaded() # What JS initially sent for user_id

            if cookie_user_id and cookie_user_id != "None":
                # If a persistent cookie ID exists, ensure we are using it.
                if current_user_id != cookie_user_id:
                    rv_user_id.set(cookie_user_id)
                    print(f"LTM ON: User ID updated to value from cookie: {cookie_user_id}")
            else:
                # No cookie, or current ID is the initial temp one. Generate/set persistent ID.
                # Note: rv_user_id might already be a persistent one if LTM was just toggled on after being off.
                # To avoid creating a new ID if one was just set by handle_user_id_cookie_load:
                # This part needs careful thought on state transitions.
                # Safest: if no cookie, always generate new *persistent* ID and set cookie.
                new_persistent_id = str(uuid.uuid4())
                rv_user_id.set(new_persistent_id)
                session.send_custom_message("setCookie", {"name": "user_id", "value": new_persistent_id, "days": 365})
                print(f"LTM ON: No user_id cookie or was temp. New persistent user_id set and cookie sent: {new_persistent_id}")

            # Load data for the (now confirmed) persistent user ID
            load_user_data_from_hf_shiny(rv_user_id.get(), rv_chat_metadata, rv_all_chat_messages)
        else: # LTM is OFF
            # Ensure user_id is temporary and no cookie is set for user_id
            rv_user_id.set(str(uuid.uuid4()))
            session.send_custom_message("deleteCookie", {"name": "user_id"}) # Ensure no orphaned persistent cookie
            print(f"LTM OFF: User ID is temporary: {rv_user_id.get()}. Any persistent user_id cookie deleted.")
            # Clear any loaded data if LTM was just turned off
            rv_chat_metadata.set({}); rv_all_chat_messages.set({})
            _initial_chat_selected_after_load.set(False) # Allow new temp session to be created

    @reactive.Effect
    @reactive.event(_cookie_ltm_processed, _cookie_user_id_processed)
    def initial_agent_and_chat_setup_trigger():
        """Sequences agent initialization and initial chat selection after cookies are processed."""
        if not (_cookie_ltm_processed.get() and _cookie_user_id_processed.get()): return
        
        print("Cookie states finalized. Proceeding with agent initialization.")
        if llm_settings_initialized:
            s_count = input.search_results_count() if input.search_results_count() is not None else 10
            agent, err = setup_agent_shiny(rv_uploaded_documents, rv_uploaded_dataframes, s_count)
            orchestrator_agent_rv.set(agent); agent_init_error_rv.set(err)
            if err: print(f"Agent init error: {err}")
            else: print("Agent initialized.")
        else:
            agent_init_error_rv.set("Agent setup skipped: LLM settings failed.")
            print(agent_init_error_rv.get())
        
        # This will now trigger select_initial_chat_after_load if not already done
        if not _initial_chat_selected_after_load.get():
            print("Triggering initial chat selection after agent setup.")
            select_initial_chat_after_load()


    @reactive.Effect
    @reactive.event(rv_all_chat_messages, rv_chat_metadata, ignore_init=True) # Observe changes from data loading
    def handle_initial_chat_selection_after_load():
        """Selects or creates a chat session after user data might have been loaded from HF."""
        # This runs if rv_all_chat_messages or rv_chat_metadata change,
        # typically after load_user_data_from_hf_shiny updates them.
        if _cookie_ltm_processed.get() and _cookie_user_id_processed.get() and not _initial_chat_selected_after_load.get():
            print("Chat data (metadata or messages) changed, attempting initial selection.")
            select_initial_chat_after_load()

    def create_new_chat_session_shiny(force_temp_name: bool = False):
        """Creates a new chat session, updating all relevant reactive values."""
        new_id = str(uuid.uuid4()); meta = rv_chat_metadata.get().copy(); name = "Current Session"
        if rv_long_term_memory_enabled.get() and not force_temp_name:
            nums = [int(m.group(1)) for n_val in meta.values() if (m := re.match(r"Idea (\d+)", n_val))]
            name = f"Idea {max(nums) + 1 if nums else 1}"
        
        meta[new_id] = name; rv_chat_metadata.set(meta)
        
        greeting = agent_generate_llm_greeting() if llm_settings_initialized and orchestrator_agent_rv.get() else "Hello! How can I assist you today? (Agent not fully ready)"
        init_msg = {"role": "assistant", "content": greeting}

        all_msgs = rv_all_chat_messages.get().copy(); all_msgs[new_id] = [init_msg]; rv_all_chat_messages.set(all_msgs)
        rv_current_chat_id.set(new_id); rv_messages.set([init_msg]); rv_chat_modified.set(False)
        rv_suggested_prompts.set(generate_suggested_prompts_shiny([init_msg]))
        print(f"New chat session created: '{name}' (ID: {new_id})")

        if rv_long_term_memory_enabled.get() and not force_temp_name:
            # Persist metadata for the new chat
            save_chat_metadata_shiny(rv_user_id.get(), meta)
            # Persist this initial message for the new chat
            save_chat_history_shiny(user_id=rv_user_id.get(), updated_chat_id=new_id, updated_messages_list=[init_msg])


    def switch_chat_shiny(chat_id: str):
        """Switches to an existing chat session or creates a new temporary one if LTM is off."""
        if not rv_long_term_memory_enabled.get():
            create_new_chat_session_shiny(force_temp_name=True); return
        if chat_id not in rv_chat_metadata.get() or rv_current_chat_id.get() == chat_id: return

        rv_current_chat_id.set(chat_id)
        # Messages for the switched-to chat should be in rv_all_chat_messages (loaded from HF or from current session)
        msgs = rv_all_chat_messages.get().get(chat_id, [])
        if not msgs: # Should not happen if data loaded correctly, but as a fallback
            print(f"Warning: No messages found in rv_all_chat_messages for chat {chat_id}. Displaying empty chat.")
            msgs = [{"role":"assistant", "content":"This chat is empty or messages could not be loaded."}]
        rv_messages.set(msgs)
        rv_suggested_prompts.set(generate_suggested_prompts_shiny(msgs))
        rv_chat_modified.set(False) # Switched to an existing chat, not modified by this action
        print(f"Switched to chat: '{rv_chat_metadata.get().get(chat_id, 'Unknown')}'")

    @reactive.Effect @reactive.event(input.new_chat_btn)
    def _(): create_new_chat_session_shiny()
    
    def process_and_send_query(query_text: str, is_regenerate: bool = False):
        """Handles sending a query, getting a response, and updating chat state."""
        if not query_text.strip(): return
        current_messages = rv_messages.get().copy()
        
        # Determine the actual query for the agent (especially for regeneration)
        actual_query_for_agent = query_text
        history_for_agent = current_messages.copy() # Start with current messages for history

        if not is_regenerate:
            current_messages.append({"role": "user", "content": query_text})
            # history_for_agent is already current_messages before assistant replies
        else: # Regeneration
            # query_text is the original user query that led to the response we're regenerating
            # We need to ensure history_for_agent doesn't include the assistant message we are about to replace.
            if history_for_agent and history_for_agent[-1]["role"] == "assistant":
                history_for_agent.pop() # Remove last assistant message before sending to agent

        rv_messages.set(current_messages) # Update UI with user message immediately (if not regenerate)

        formatted_hist = format_chat_history_shiny(history_for_agent)
        temp = rv_llm_temperature.get(); verbosity = rv_llm_verbosity.get(); search_n = rv_search_results_count.get()
        assistant_response = get_agent_response_shiny(actual_query_for_agent, formatted_hist, temp, verbosity, search_n)

        # final_msgs_for_rv is based on the messages list *after* user message might have been added
        final_msgs_for_rv = rv_messages.get().copy()
        if is_regenerate and final_msgs_for_rv and final_msgs_for_rv[-1]["role"] == "assistant":
            final_msgs_for_rv.pop() # Remove the old assistant message that is being regenerated

        final_msgs_for_rv.append({"role": "assistant", "content": assistant_response})
        rv_messages.set(final_msgs_for_rv)

        rv_suggested_prompts.set(generate_suggested_prompts_shiny(final_msgs_for_rv))
        rv_chat_modified.set(True)

        if rv_long_term_memory_enabled.get():
            uid = rv_user_id.get(); cid = rv_current_chat_id.get()
            all_msgs = rv_all_chat_messages.get().copy()
            all_msgs[cid] = final_msgs_for_rv # Update the master list of all messages
            rv_all_chat_messages.set(all_msgs)
            if uid and cid: save_chat_history_shiny(user_id=uid, updated_chat_id=cid, updated_messages_list=final_msgs_for_rv)

    @reactive.Effect @reactive.event(input.send_message_btn)
    def _(): query = input.chat_input(); process_and_send_query(query); ui.update_text_area("chat_input", value="")

    @reactive.Effect @reactive.event(input.file_uploader)
    def _():
        infos = input.file_uploader(); msgs_to_add = []
        for fi in infos if infos else []:
            ft, name = process_uploaded_file_shiny(fi, rv_uploaded_documents, rv_uploaded_dataframes)
            msg = f"Doc '{name}' processed." if ft=="document" else f"Dataset '{name}' processed." if ft=="dataframe" else f"File '{name}' not processed."
            msgs_to_add.append({"role":"assistant", "content":msg})
        if msgs_to_add: cur = rv_messages.get().copy(); cur.extend(msgs_to_add); rv_messages.set(cur); rv_suggested_prompts.set(generate_suggested_prompts_shiny(cur))

    def remove_file_shiny(file_name: str, type_to_remove: str):
        msg = f"File '{file_name}' "; target_rv = rv_uploaded_documents if type_to_remove == "document" else rv_uploaded_dataframes
        data = target_rv.get().copy()
        if file_name in data: del data[file_name]; target_rv.set(data); msg += f"removed from {type_to_remove}s."
        else: msg += f"not found in {type_to_remove}s."
        ws_path = os.path.join(UI_ACCESSIBLE_WORKSPACE, file_name)
        try:
            if os.path.exists(ws_path): os.remove(ws_path); msg += " Also deleted from workspace."
        except Exception as e: msg += f" Error deleting from workspace: {e}"
        cur = rv_messages.get().copy(); cur.append({"role": "assistant", "content": msg}); rv_messages.set(cur)

    # LLM Settings Effects
    @reactive.Effect @reactive.event(input.llm_temperature)
    def _(): rv_llm_temperature.set(input.llm_temperature()); # print(f"Temp set to {rv_llm_temperature.get()}") # Reduced logging
    @reactive.Effect @reactive.event(input.llm_verbosity)
    def _(): rv_llm_verbosity.set(input.llm_verbosity()); # print(f"Verbosity set to {rv_llm_verbosity.get()}") # Reduced logging
    
    @reactive.Effect @reactive.event(input.search_results_count, ignore_init=True)
    def handle_search_results_change_reinit_agent(): # More descriptive name
        val = input.search_results_count(); rv_search_results_count.set(val)
        # print(f"Search results count changed to: {val}. Re-initializing agent.") # Reduced logging
        if llm_settings_initialized:
            agent, err = setup_agent_shiny(rv_uploaded_documents, rv_uploaded_dataframes, val)
            orchestrator_agent_rv.set(agent); agent_init_error_rv.set(err)
            if err: print(f"Agent re-init error: {err}") # Keep error logging
        else: agent_init_error_rv.set("Agent re-init skipped: LLM settings failed.")

    # Clipboard Effects
    @reactive.Effect @reactive.event(rv_clipboard_text)
    def _(): text = rv_clipboard_text.get(); session.send_custom_message("copyToClipboard", {"text": text}) if text else None; rv_clipboard_text.set(None)
    @reactive.Effect @reactive.event(input.clipboard_copy_success)
    def _(): ui.notification_show("Copied!", duration=3, type="default", session=session)
    @reactive.Effect @reactive.event(input.clipboard_copy_error)
    def _(): ui.notification_show(f"Copy failed: {input.clipboard_copy_error().get('error','')}", duration=5, type="error", session=session)
    
    @reactive.Effect
    def handle_dynamic_action_buttons():
        """
        Centralized handler for dynamically generated action buttons.
        It compares current input values of known button prefixes against their last known values.
        """
        current_values = {k: v for k, v in session.input.to_dict().items() if isinstance(v, int)}
        last_values = last_action_button_values.get()
        clicked_button_id = None # Store ID of button that was actually clicked

        for k, v in current_values.items():
            if v > last_values.get(k, 0): # A click means the value increments
                clicked_button_id = k # This is the button that was clicked
                # print(f"Dynamic button clicked: {k}") # Can be noisy, removed for now
                id_parts = k.split("_")
                action_type = id_parts[0] # e.g., "remove", "copy", "select", "suggested", "options"

                # Dispatch based on prefix more robustly
                if k.startswith("remove_doc_"): remove_file_shiny("_".join(id_parts[2:]), "document") # Name might have underscores
                elif k.startswith("remove_df_"): remove_file_shiny("_".join(id_parts[2:]), "dataframe")
                elif k.startswith("copy_msg_"): idx=int(id_parts[2]); rv_clipboard_text.set(rv_messages.get()[idx]['content'])
                elif k.startswith("regenerate_msg_"):
                    idx=int(id_parts[2]); msgs_list=rv_messages.get()
                    if idx == 0 and msgs_list[idx]["role"] == "assistant": # Regenerate greeting
                        new_greeting = agent_generate_llm_greeting() if llm_settings_initialized and orchestrator_agent_rv.get() else "Hello (fallback greeting)!"
                        updated_msgs = msgs_list.copy(); updated_msgs[0]['content'] = new_greeting; rv_messages.set(updated_msgs)
                        rv_suggested_prompts.set(generate_suggested_prompts_shiny(updated_msgs))
                        if rv_long_term_memory_enabled.get(): save_chat_history_shiny(rv_user_id.get(),rv_current_chat_id.get(),updated_msgs)
                    elif idx > 0 and msgs_list[idx-1]["role"] == "user" and msgs_list[idx]["role"] == "assistant":
                        process_and_send_query(msgs_list[idx-1]['content'], is_regenerate=True)
                elif k.startswith("select_chat_"): switch_chat_shiny("_".join(id_parts[2:])) # Chat ID might have underscores
                elif k.startswith("suggested_prompt_"): idx=int(id_parts[2]); process_and_send_query(DEFAULT_PROMPTS[idx]) # Use DEFAULT_PROMPTS as source
                elif k.startswith("options_chat_"): rv_options_modal_chat_id.set("_".join(id_parts[2:]))
                break # Process one click at a time
        
        if clicked_button_id: # Update last_values only for the clicked button
            new_last_values = last_values.copy()
            new_last_values[clicked_button_id] = current_values[clicked_button_id]
            last_action_button_values.set(new_last_values)
        # For other buttons, their state doesn't change in last_values unless clicked.
        # This prevents re-triggering if multiple button states are sent by Shiny at once
        # but only one was actually incremented by a click.

    # Modal Effects
    @reactive.Effect @reactive.event(input.modal_save_rename_btn)
    def _():
        chat_id=rv_options_modal_chat_id.get();new_name=input.modal_rename_chat_input()
        if chat_id and new_name and new_name.strip():
            meta=rv_chat_metadata.get().copy();meta[chat_id]=new_name.strip();rv_chat_metadata.set(meta)
            if rv_long_term_memory_enabled.get(): save_chat_metadata_shiny(rv_user_id.get(), meta)
        ui.modal_remove()
    @reactive.Effect @reactive.event(input.modal_delete_chat_btn)
    def _(): ui.modal_remove(); rv_show_delete_confirmation_modal.set(True)
    @reactive.Effect
    def _():
        chat_id=rv_options_modal_chat_id.get()
        if chat_id and not rv_show_delete_confirmation_modal.get(): ui.modal_show(chat_options_modal_ui(chat_id, rv_chat_metadata.get().get(chat_id,"Chat")))
    @reactive.Effect
    def _():
        if rv_show_delete_confirmation_modal.get():
            chat_id=rv_options_modal_chat_id.get()
            if chat_id: ui.modal_show(delete_chat_confirmation_modal_ui(rv_chat_metadata.get().get(chat_id,"Selected Chat")))
    @reactive.Effect @reactive.event(input.modal_confirm_delete_btn)
    def _():
        chat_id_to_delete=rv_options_modal_chat_id.get()
        if chat_id_to_delete:
            meta=rv_chat_metadata.get().copy();del meta[chat_id_to_delete];rv_chat_metadata.set(meta)
            all_m=rv_all_chat_messages.get().copy();all_m.pop(chat_id_to_delete,None);rv_all_chat_messages.set(all_m)
            if rv_long_term_memory_enabled.get():
                current_uid = rv_user_id.get()
                if current_uid:
                    save_chat_metadata_shiny(current_uid, meta)
                    save_chat_history_shiny(user_id=current_uid, messages_dict_to_save_directly=all_m)
            if rv_current_chat_id.get()==chat_id_to_delete:
                if meta: switch_chat_shiny(next(iter(sorted(meta.keys())), None)) # Switch to first available or None
                else: create_new_chat_session_shiny()
        ui.modal_remove();rv_show_delete_confirmation_modal.set(False);rv_options_modal_chat_id.set(None)
    @reactive.Effect @reactive.event(input.modal_cancel_delete_btn)
    def _(): ui.modal_remove();rv_show_delete_confirmation_modal.set(False);rv_options_modal_chat_id.set(None)

    # Output Renderers (using .get() for all reactive value access)
    @output @render.ui
    def chat_history_output():
        items = [ui.input_action_button("new_chat_btn",f"➕ {'Temp ' if not rv_long_term_memory_enabled.get() else ''}New Chat",class_="btn-primary w-100 mb-2")]
        if not rv_long_term_memory_enabled.get(): items.append(ui.tags.div(ui.tags.strong("Warning:")," LTM disabled.",style="color:orange;font-size:0.9em;"))
        if rv_long_term_memory_enabled.get():
            meta = rv_chat_metadata.get()
            if not meta: items.append(ui.tags.p("No saved chats.",style="text-align:center;color:grey;"))
            else:
                for cid, name in sorted(meta.items(), key=lambda x: x[1]): # Sort by name
                    is_curr = cid == rv_current_chat_id.get()
                    items.append(ui.tags.div(ui.row(ui.column(9,ui.input_action_button(f"select_chat_{cid}",name,class_=f"{'btn-primary' if is_curr else 'btn-light'} w-100 text-start")),ui.column(3,ui.input_action_button(f"options_chat_{cid}","⚙️",class_="btn-secondary"))),style="margin-bottom:5px;"))
        return ui.TagList(*items)

    @output @render.ui
    def uploaded_files_output():
        items=[];docs=rv_uploaded_documents.get();dfs=rv_uploaded_dataframes.get()
        if not docs and not dfs: items.append(ui.tags.p("No files.",style="text-align:center;color:grey;"))
        if docs:
            items.append(ui.tags.h6("Documents:",style="margin-top:10px"))
            for i, name in enumerate(docs.keys()): items.append(ui.tags.div(ui.row(ui.column(1,ui.tags.span("📄")),ui.column(8,ui.tags.span(name)),ui.column(3,ui.input_action_button(f"remove_doc_{i}_{name}","🗑️",class_="btn-sm btn-danger")))))
        if dfs:
            items.append(ui.tags.h6("Datasets:",style="margin-top:10px"))
            for i, name in enumerate(dfs.keys()): items.append(ui.tags.div(ui.row(ui.column(1,ui.tags.span("📊")),ui.column(8,ui.tags.span(name)),ui.column(3,ui.input_action_button(f"remove_df_{i}_{name}","🗑️",class_="btn-sm btn-danger")))))
        return ui.TagList(*items)

    @output @render.ui
    def chat_display_output():
        msgs=rv_messages.get()
        if not msgs: return ui.tags.div("No messages yet. Start by typing below or use a suggestion!",style="text-align:center;color:grey;padding:20px;min-height:300px;")
        elements = []
        for i,m in enumerate(msgs):
            r,c = m.get("role","unknown"),m.get("content","")
            is_u = r=="user"; style_str = f"background-color:{'#DCF8C6' if is_u else '#E0E0E0'};margin-left:{'auto' if is_u else '10px'};margin-right:{'10px' if is_u else 'auto'};padding:10px;border-radius:15px;max-width:70%;word-wrap:break-word;margin-bottom:5px;"
            btns_ui = [ui.input_action_button(f"copy_msg_{i}","📄",class_="btn-xs btn-outline-secondary")]
            if r=="assistant" and i==len(msgs)-1: btns_ui.append(ui.input_action_button(f"regenerate_msg_{i}","🔄",class_="btn-xs btn-outline-warning"))
            elements.append(ui.tags.div(ui.tags.div(ui.tags.strong(f"{r.capitalize()}:")),ui.markdown(c),ui.tags.div(*btns_ui,style="margin-top:5px;text-align:right;"),style=style_str,class_=f"message-bubble {'user' if is_u else 'assistant'}-bubble"))
        return ui.TagList(*elements,style="display:flex;flex-direction:column;padding:10px;min-height:300px;")

    @output @render.ui
    def suggested_prompts_output():
        prompts=rv_suggested_prompts.get()
        if not prompts: return ui.tags.div()
        btns = [ui.input_action_button(f"suggested_prompt_{i}",p,class_="btn-sm btn-outline-secondary m-1") for i,p in enumerate(prompts)]
        return ui.TagList(ui.tags.h5("Suggestions:",style="margin-bottom:5px;"),ui.tags.div(*btns,style="display:flex;flex-wrap:wrap;justify-content:center;"))

    @output @render.text
    def startup_status():
        s=[];agent_e=agent_init_error_rv.get();agent_i=orchestrator_agent_rv.get()
        if llm_settings_error:s.append(f"LLM Err:{llm_settings_error}")
        elif llm_settings_initialized:s.append("LLM OK.")
        if agent_e:s.append(f"Agent Err:{agent_e}")
        elif agent_i:s.append("Agent OK.")
        if not os.getenv("GOOGLE_API_KEY"):s.append("GOOG_KEY Missing.")
        # print(f"UI Startup Status: {' | '.join(s) if s else 'System Ready.'}") # Reduced console logging
        return " | ".join(s) if s else "System Ready."
    
    print(f"--- Server Function Initialized for session {session.id} ---")


app = App(app_ui, server)
